#!/usr/bin/env python3
"""Generate comprehensive project documentation as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

# ── Title Page ──
doc.add_paragraph('')
doc.add_paragraph('')
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Agent Skills for All')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Engineering Documentation')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Repository: ').bold = True
meta.add_run('github.com/harishrajora/agentskillsforall\n')
meta.add_run('CLI Package: ').bold = True
meta.add_run('npx agentskillshub-cli\n')
meta.add_run('Website: ').bold = True
meta.add_run('https://www.agentskillsforall.com\n')
meta.add_run('npm Package: ').bold = True
meta.add_run('agentskillshub-cli')

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. System Architecture & Flowchart',
    '3. Repository Structure',
    '4. CLI Package (skills/)',
    '   4.1 How the CLI Works',
    '   4.2 CLI Entry Point & Command Routing',
    '   4.3 Key Source Files',
    '   4.4 Skill Installation Flow',
    '   4.5 Telemetry System',
    '   4.6 Agents Support',
    '   4.7 Publishing to npm',
    '5. Website (website/)',
    '   5.1 Tech Stack',
    '   5.2 Pages & Routing',
    '   5.3 API Endpoints',
    '   5.4 Component Architecture',
    '   5.5 Library Files',
    '   5.6 Skills Data & Catalog',
    '   5.7 Telemetry-Driven Discovery',
    '6. Data Flow: End-to-End',
    '7. Environment Variables',
    '8. Deployment Guide',
    '9. Development Workflow',
    '10. Key Design Decisions',
    '11. Troubleshooting',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ══════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'Agent Skills for All is an open ecosystem for discovering, installing, and managing '
    'reusable "skills" for AI coding agents (Claude Code, Cursor, GitHub Copilot, Gemini CLI, '
    'Codex, and 40+ others). A "skill" is a SKILL.md file — a markdown document with frontmatter '
    'that gives an AI agent specialized knowledge about a framework, tool, or workflow.'
)

doc.add_paragraph(
    'The project consists of two main components that live in a single monorepo:'
)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
for i, (comp, path, desc) in enumerate([
    ('Component', 'Directory', 'Description'),
    ('CLI', 'skills/', 'npm package "agentskillshub-cli" — installs skills from GitHub repos into agent directories'),
    ('Website', 'website/', 'Next.js app at agentskillsforall.com — skill catalog, search API, telemetry'),
]):
    for j, val in enumerate([comp, path, desc]):
        cell = table.cell(i, j)
        cell.text = val
        if i == 0:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph('')
doc.add_paragraph(
    'When a user runs "npx agentskillshub-cli add <github-url>", the CLI clones the repo, '
    'discovers SKILL.md files, and symlinks/copies them into the appropriate agent directory '
    '(e.g., .claude/skills/, .cursor/skills/). It also sends a telemetry event to the website, '
    'which tracks install counts and makes skills discoverable via search.'
)

# ══════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE & FLOWCHART
# ══════════════════════════════════════════════
doc.add_heading('2. System Architecture & Flowchart', level=1)

doc.add_heading('2.1 High-Level Architecture', level=2)

flowchart = """
┌─────────────────────────────────────────────────────────────────────┐
│                        USER'S TERMINAL                              │
│                                                                     │
│   $ npx agentskillshub-cli add https://github.com/owner/repo       │
└────────────────────────────┬────────────────────────────────────────┘
                             │
                             ▼
┌─────────────────────────────────────────────────────────────────────┐
│                     CLI (agentskillshub-cli)                        │
│                                                                     │
│  1. Parse source URL                                                │
│  2. Try blob install (fast path) from website API                   │
│  3. If blob fails → git clone the repository                        │
│  4. Discover SKILL.md files in repo                                 │
│  5. Show interactive selection (if multiple skills)                  │
│  6. Detect installed agents (Claude, Cursor, Copilot, etc.)         │
│  7. Symlink/copy skill files to agent directories                   │
│  8. Send telemetry event to website API ──────────────────────┐     │
│  9. Update local lock file (.skill-lock.json)                 │     │
└───────────────────────────────────────────────────────────────┘     │
                                                                      │
                             ┌────────────────────────────────────────┘
                             ▼
┌─────────────────────────────────────────────────────────────────────┐
│                    WEBSITE (agentskillsforall.com)                   │
│                                                                     │
│  ┌──────────────┐  ┌──────────────┐  ┌────────────────────┐        │
│  │  /api/t      │  │ /api/search  │  │ /api/github-stats  │        │
│  │  (telemetry) │  │ (skill find) │  │ (stars/forks)      │        │
│  └──────┬───────┘  └──────┬───────┘  └────────────────────┘        │
│         │                 │                                         │
│         ▼                 ▼                                         │
│  ┌──────────────────────────────┐                                   │
│  │     In-Memory Store          │                                   │
│  │  (skills map + event log)    │                                   │
│  └──────────────────────────────┘                                   │
│         │                                                           │
│         ▼                                                           │
│  ┌──────────────────────────────┐    ┌─────────────────────┐        │
│  │   Homepage (SSR)             │    │  /skills/[slug]     │        │
│  │   - Static catalog (45)     │◄───│  Skill detail page  │        │
│  │   - + telemetry skills      │    │  - Fetches SKILL.md │        │
│  │   - Leaderboard table       │    │    from GitHub      │        │
│  └──────────────────────────────┘    └─────────────────────┘        │
└─────────────────────────────────────────────────────────────────────┘

                             ▲
                             │
┌─────────────────────────────────────────────────────────────────────┐
│                        GITHUB                                       │
│                                                                     │
│  owner/repo                                                         │
│  ├── SKILL.md              (skill definition)                       │
│  ├── skills/                                                        │
│  │   ├── skill-a/SKILL.md                                           │
│  │   └── skill-b/SKILL.md                                           │
│  └── README.md                                                      │
└─────────────────────────────────────────────────────────────────────┘
"""

p = doc.add_paragraph()
run = p.add_run(flowchart)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_heading('2.2 Skill Installation Flow', level=2)

install_flow = """
User runs: npx agentskillshub-cli add <source>
    │
    ├─► Parse source (GitHub URL, owner/repo, local path, git URL)
    │
    ├─► Is it a GitHub owner/repo format?
    │     ├─ YES ─► Try blob install (fast path via /api/download)
    │     │           ├─ Success ─► Skip clone, use cached files
    │     │           └─ Fail ─► Fall back to git clone
    │     └─ NO ──► Git clone directly
    │
    ├─► Discover SKILL.md files in cloned/downloaded repo
    │     └─ Searches: root, skills/, .agents/skills/, .claude/skills/, etc.
    │
    ├─► Filter skills (if --skill flag provided)
    │     └─ If no filter + multiple skills ─► Show interactive selector
    │
    ├─► Detect installed agents on the machine
    │     └─ Checks for: .claude/, .cursor/, .github/copilot, etc.
    │
    ├─► Install skills to agent directories
    │     ├─ Symlink (default) or copy (--copy flag)
    │     └─ Creates: <agent-dir>/skills/<skill-name>/SKILL.md
    │
    ├─► Security audit (fetch from /api/audit)
    │
    ├─► Send telemetry to /api/t
    │     └─ event=install, source=owner/repo, skills=skill-name
    │
    └─► Update lock file (.skill-lock.json)
"""

p = doc.add_paragraph()
run = p.add_run(install_flow)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('2.3 Website Data Flow', level=2)

web_flow = """
Homepage Request (GET /)
    │
    ├─► Server Component: getSkills()
    │     ├─► Read skills/index.json (45 pre-loaded skills)
    │     ├─► Fetch /api/search (telemetry-installed skills)
    │     ├─► Merge both lists
    │     └─► Sort: Hot skills first, then alphabetical
    │
    ├─► Render SkillsList component
    │     └─ Client-side search + category filters
    │
    └─► User clicks a skill
          │
          ▼
        Skill Detail Page (GET /skills/[slug])
          │
          ├─► Try local markdown files (skills/<path>/SKILL.md)
          ├─► If not found + has sourceRepo ─► Fetch from GitHub
          │     ├─ Try: <skill>/SKILL.md
          │     ├─ Try: SKILL.md (root)
          │     └─ Fallback: README.md
          │
          └─► Render with ReactMarkdown + rehype-raw + rehype-sanitize
"""

p = doc.add_paragraph()
run = p.add_run(web_flow)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# ══════════════════════════════════════════════
# 3. REPOSITORY STRUCTURE
# ══════════════════════════════════════════════
doc.add_heading('3. Repository Structure', level=1)

structure = """
agentskillsforall/
├── package.json              # Root package.json (for Vercel Next.js detection)
├── vercel.json               # Vercel deployment config
├── .gitignore                # Git ignore rules
│
├── skills/                   # CLI PACKAGE (npm: agentskillshub-cli)
│   ├── package.json          # CLI package config (name: agentskillshub-cli)
│   ├── bin/
│   │   └── cli.mjs           # Entry point (#!/usr/bin/env node)
│   ├── src/
│   │   ├── cli.ts            # Main CLI logic, command routing, ASCII logo
│   │   ├── add.ts            # "add" command — skill installation
│   │   ├── find.ts           # "find" command — interactive search
│   │   ├── list.ts           # "list" command — show installed skills
│   │   ├── remove.ts         # "remove" command — uninstall skills
│   │   ├── install.ts        # Lock file based installation
│   │   ├── sync.ts           # Sync from node_modules
│   │   ├── installer.ts      # Core install logic (symlink/copy)
│   │   ├── agents.ts         # 44 agent definitions and paths
│   │   ├── blob.ts           # Fast blob-based install from API
│   │   ├── telemetry.ts      # Telemetry event sending
│   │   ├── source-parser.ts  # Parse GitHub URLs, owner/repo, etc.
│   │   ├── skill-lock.ts     # Global lock file management
│   │   ├── local-lock.ts     # Project-level lock file
│   │   ├── update-source.ts  # Build install URLs for updates
│   │   ├── frontmatter.ts    # YAML frontmatter parser
│   │   ├── git.ts            # Git operations (clone, etc.)
│   │   ├── types.ts          # TypeScript interfaces
│   │   └── constants.ts      # Shared constants
│   ├── skills/
│   │   └── find-skills/
│   │       └── SKILL.md      # Built-in "find skills" skill
│   ├── tests/                # Test files
│   ├── scripts/              # Build & utility scripts
│   └── dist/                 # Compiled output (git-ignored)
│
├── website/                  # NEXT.JS WEBSITE
│   ├── package.json          # Website dependencies
│   ├── next.config.ts        # Next.js configuration
│   ├── tsconfig.json         # TypeScript config (paths: @/* → ./src/*)
│   ├── src/
│   │   ├── app/
│   │   │   ├── layout.tsx    # Root layout (dark theme, fonts, header, footer)
│   │   │   ├── page.tsx      # Homepage (hero + skills listing)
│   │   │   ├── globals.css   # Global styles (dark theme, animations)
│   │   │   ├── stripe-dev-theme.css  # Hero section styles
│   │   │   ├── skills/
│   │   │   │   └── [slug]/
│   │   │   │       ├── layout.tsx          # Skill header, breadcrumb, install cmd
│   │   │   │       ├── page.tsx            # Overview tab
│   │   │   │       ├── not-found.tsx       # 404 for skills
│   │   │   │       ├── documentation/
│   │   │   │       │   └── page.tsx        # Documentation tab
│   │   │   │       └── [docType]/
│   │   │   │           └── page.tsx        # Playbook/Advanced/Cloud tabs
│   │   │   └── api/
│   │   │       ├── search/route.ts         # GET /api/search — skill search
│   │   │       ├── t/route.ts              # GET /api/t — telemetry ingestion
│   │   │       ├── github-stats/route.ts   # GET /api/github-stats
│   │   │       ├── audit/route.ts          # GET /api/audit — stub
│   │   │       └── download/[owner]/[repo]/[slug]/route.ts  # Download stub
│   │   ├── components/
│   │   │   ├── Header.tsx             # Navbar with logo
│   │   │   ├── HeroTitle.tsx          # Parallax hero title
│   │   │   ├── HeroCtaRow.tsx         # CTA buttons
│   │   │   ├── AgentsMarquee.tsx      # Scrolling agent logos
│   │   │   ├── InstallCommandBlock.tsx # Copy-to-clipboard install command
│   │   │   ├── SkillsList.tsx         # Leaderboard table with search/filters
│   │   │   ├── SkillCard.tsx          # Individual skill row (unused now)
│   │   │   ├── SkillLogo.tsx          # Skill framework icon
│   │   │   ├── SkillDocNav.tsx        # Tab navigation for skill docs
│   │   │   ├── SkillDocView.tsx       # Markdown renderer
│   │   │   ├── SkillOverview.tsx      # Parsed skill description
│   │   │   ├── SkillOverviewDocView.tsx # Enhanced overview markdown
│   │   │   ├── CopyableUrlBlock.tsx   # URL/command copy block
│   │   │   ├── GitHubStatsInline.tsx  # Stars/forks per skill repo
│   │   │   ├── NavbarGitHubStats.tsx  # Animated stats in navbar
│   │   │   ├── NavLinks.tsx           # Navigation links
│   │   │   ├── SectionLabel.tsx       # Section divider
│   │   │   ├── ScrollReveal.tsx       # Scroll animation observer
│   │   │   └── ui/                    # shadcn/ui primitives
│   │   │       ├── button.tsx
│   │   │       ├── badge.tsx
│   │   │       ├── card.tsx
│   │   │       ├── input.tsx
│   │   │       ├── select.tsx
│   │   │       └── tabs.tsx
│   │   └── lib/
│   │       ├── skills.ts          # Core skill loading & merging
│   │       ├── skillsFs.ts        # Filesystem skill reading (server-only)
│   │       ├── skillsHelpers.ts   # Display names, hot skills, sorting
│   │       ├── skillsMapping.ts   # Path → LambdaTest name mapping
│   │       ├── skillIcons.ts      # Skill icon URL mapping
│   │       ├── install.ts         # Agent paths & git clone commands
│   │       ├── store.ts           # In-memory telemetry store
│   │       ├── github.ts          # GitHub API (stars/forks)
│   │       ├── types.ts           # TypeScript interfaces
│   │       ├── parseDescription.ts     # Parse skill descriptions
│   │       ├── parseSkillMarkdown.ts   # Split markdown by headings
│   │       ├── proseClasses.ts    # Tailwind prose classes
│   │       └── utils.ts           # cn() utility (clsx + tailwind-merge)
│   ├── skills/                    # 45 pre-loaded skill directories
│   │   ├── index.json             # Skills catalog metadata
│   │   ├── selenium-automation-skill/
│   │   ├── playwright-automation-skill/
│   │   ├── cypress-automation-skill/
│   │   └── ... (45 total)
│   └── public/
│       ├── favicon.ico
│       ├── skill-icons/           # 47 SVG framework icons
│       └── agents/                # 15 SVG agent logos
"""

p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Courier New'
run.font.size = Pt(8)

doc.add_page_break()

# ══════════════════════════════════════════════
# 4. CLI PACKAGE
# ══════════════════════════════════════════════
doc.add_heading('4. CLI Package (skills/)', level=1)

doc.add_heading('4.1 How the CLI Works', level=2)
doc.add_paragraph(
    'The CLI is published to npm as "agentskillshub-cli". When a user runs '
    '"npx agentskillshub-cli", npm downloads the package and executes bin/cli.mjs, '
    'which imports the compiled dist/cli.mjs. The source is written in TypeScript '
    'and compiled using obuild (a lightweight bundler).'
)
doc.add_paragraph(
    'The CLI supports these commands:'
)

cmds = [
    ('add <source>', 'Install skills from a GitHub repo, URL, or local path'),
    ('find [query]', 'Interactive search for skills via the website API'),
    ('list / ls', 'List installed skills (project or global)'),
    ('remove / rm', 'Remove installed skills from agent directories'),
    ('update / upgrade', 'Update installed skills to latest versions'),
    ('init [name]', 'Create a new SKILL.md template'),
    ('experimental_install', 'Restore skills from skills-lock.json'),
    ('experimental_sync', 'Sync skills from node_modules'),
]
table = doc.add_table(rows=len(cmds)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Command'
table.cell(0, 1).text = 'Description'
table.cell(0, 0).paragraphs[0].runs[0].bold = True
table.cell(0, 1).paragraphs[0].runs[0].bold = True
for i, (cmd, desc) in enumerate(cmds):
    table.cell(i+1, 0).text = cmd
    table.cell(i+1, 1).text = desc

doc.add_heading('4.2 CLI Entry Point & Command Routing', level=2)
doc.add_paragraph(
    'bin/cli.mjs is the executable entry point. It enables Node.js compile cache '
    'for performance, then imports dist/cli.mjs (the compiled version of src/cli.ts).'
)
doc.add_paragraph(
    'src/cli.ts contains the main() function which:'
)
items = [
    'Parses process.argv to determine the command and flags',
    'Shows the ASCII art logo ("AGENT SKILLS FOR ALL" in block letters)',
    'Routes to the appropriate command handler (runAdd, runFind, runList, etc.)',
    'Handles --help, --version, and unknown commands',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 Key Source Files', level=2)

files = [
    ('src/cli.ts', '920 lines',
     'Main entry point. Contains the ASCII logo, banner display, help text, '
     'command routing (switch/case on command name), and the entire update/check flow. '
     'The showBanner() function displays the logo + available commands. The main() function '
     'parses argv and dispatches to command handlers.'),
    ('src/add.ts', '1850+ lines',
     'The largest file — handles the "add" command. Key functions:\n'
     '• parseAddOptions() — parses CLI flags (-g, -a, -s, --copy, --all, etc.)\n'
     '• runAdd() — main entry point for skill installation\n'
     '• Tries blob install first (fast path), falls back to git clone\n'
     '• Discovers SKILL.md files, shows interactive skill selector\n'
     '• Detects installed agents, creates installation plan\n'
     '• Calls installer.ts to perform actual file operations\n'
     '• Sends telemetry and updates lock files'),
    ('src/find.ts', '350 lines',
     'Interactive skill search. Uses the website API (/api/search) to find skills. '
     'Features a custom fzf-style terminal UI with real-time search, arrow key navigation, '
     'debounced API calls, and direct installation of selected skills.'),
    ('src/agents.ts', '~500 lines',
     'Defines all 44 supported AI coding agents. Each agent has:\n'
     '• id — unique identifier (e.g., "claude-code")\n'
     '• name — display name (e.g., "Claude Code")\n'
     '• projectSkillsPath — where skills go in a project (e.g., ".claude/skills")\n'
     '• globalSkillsPath — where global skills go (e.g., "~/.claude/skills")\n'
     '• isInstalled() — detection function'),
    ('src/installer.ts', '700+ lines',
     'Core installation logic. Handles:\n'
     '• Symlinking skill files to agent directories (default)\n'
     '• Copying files when --copy flag is used\n'
     '• Blob-based installation (from pre-cached snapshots)\n'
     '• Creating .agents/skills/ directory structure'),
    ('src/blob.ts', '430 lines',
     'Fast installation path. Instead of cloning a git repo, fetches skill files '
     'from the website download API (/api/download). Steps:\n'
     '1. Fetch repo tree from GitHub Trees API\n'
     '2. Discover SKILL.md paths in the tree\n'
     '3. Fetch frontmatter from raw.githubusercontent.com\n'
     '4. Fetch full file snapshots from website API\n'
     'Only enabled for allowlisted owners (LambdaTest, harishrajora).'),
    ('src/telemetry.ts', '~150 lines',
     'Fire-and-forget telemetry. Sends install/remove/update/find events to the '
     'website API. Uses SKILLS_API_URL env var for the base URL (defaults to '
     'https://www.agentskillsforall.com). Respects DISABLE_TELEMETRY and '
     'DO_NOT_TRACK environment variables. Never blocks the CLI — errors are silently ignored.'),
    ('src/source-parser.ts', '~200 lines',
     'Parses various source formats into a normalized structure:\n'
     '• owner/repo → https://github.com/owner/repo.git\n'
     '• owner/repo@skill-name → repo + skill filter\n'
     '• owner/repo#branch → repo + ref\n'
     '• Full GitHub/GitLab URLs\n'
     '• git@host:owner/repo.git (SSH)\n'
     '• Local paths (./my-skills)'),
    ('src/skill-lock.ts', '~170 lines',
     'Manages the global lock file (~/.agents/.skill-lock.json). Tracks installed skills '
     'with source, version, folder hash, and timestamps. Used by the update command '
     'to detect which skills have newer versions available.'),
]

for fname, lines, desc in files:
    p = doc.add_paragraph()
    run = p.add_run(f'{fname} ({lines})')
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(desc)

doc.add_heading('4.4 Skill Installation Flow (Detailed)', level=2)
doc.add_paragraph(
    'When "npx agentskillshub-cli add owner/repo" is run:'
)
steps = [
    'Source Parsing: source-parser.ts converts the input to a normalized {url, ref, subpath, skillFilter} object.',
    'Blob Fast Path: For allowlisted owners, blob.ts tries to fetch skills without cloning. It uses the GitHub Trees API to discover SKILL.md files, fetches frontmatter to get names, then downloads pre-cached snapshots from /api/download.',
    'Git Clone Fallback: If blob fails, git.ts clones the repo to a temp directory. Shallow clone (--depth 1) for speed.',
    'Skill Discovery: Scans the cloned repo for SKILL.md files. Checks priority directories (root, skills/, .agents/skills/, etc.) first. Parses frontmatter for name and description.',
    'Skill Selection: If --skill flag provided, filters to matching skills. Otherwise, if multiple skills found, shows an interactive multi-select prompt using @clack/prompts.',
    'Agent Detection: agents.ts checks which coding agents are installed on the machine by looking for their config directories.',
    'Installation: installer.ts symlinks (or copies) skill files from the temp directory to each detected agent\'s skills directory. Creates .agents/skills/<name>/ as a universal location.',
    'Security Audit: Fetches audit data from /api/audit for the installed skills. Displays risk levels if available.',
    'Telemetry: Sends an install event to /api/t with source, skill names, agent names, and version.',
    'Lock File: Updates .skill-lock.json with the skill\'s source, hash, and timestamps for future update checks.',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_heading('4.5 Telemetry System', level=2)
doc.add_paragraph(
    'The CLI sends anonymous telemetry events to the website. This serves two purposes:'
)
doc.add_paragraph('1. Populates the skill search index — skills appear in "npx agentskillshub-cli find" results based on install counts.')
doc.add_paragraph('2. Powers the homepage leaderboard — newly installed skills appear on the website.')
doc.add_paragraph(
    'Telemetry is fire-and-forget (non-blocking). It can be disabled via DISABLE_TELEMETRY=1 '
    'or DO_NOT_TRACK=1 environment variables. The telemetry URL defaults to '
    'https://www.agentskillsforall.com/api/t but can be overridden with SKILLS_API_URL.'
)

doc.add_heading('4.6 Supported Agents (44 Total)', level=2)
doc.add_paragraph(
    'The CLI supports installing skills to 44 different AI coding agents. Key agents include:'
)
agents = [
    ('Claude Code', '.claude/skills/', '~/.claude/skills/'),
    ('Cursor', '.cursor/skills/', '~/.cursor/skills/'),
    ('GitHub Copilot', '.github/skills/', '~/.github/skills/'),
    ('Gemini CLI', '.gemini/skills/', '~/.gemini/skills/'),
    ('Codex', '.codex/skills/', '~/.codex/skills/'),
    ('Windsurf', '.windsurf/skills/', '~/.windsurf/skills/'),
    ('Cline', '.cline/skills/', '~/.cline/skills/'),
    ('OpenCode', '.opencode/skills/', '~/.opencode/skills/'),
]
table = doc.add_table(rows=len(agents)+1, cols=3)
table.style = 'Light Grid Accent 1'
for j, h in enumerate(['Agent', 'Project Path', 'Global Path']):
    table.cell(0, j).text = h
    table.cell(0, j).paragraphs[0].runs[0].bold = True
for i, (name, proj, glob) in enumerate(agents):
    table.cell(i+1, 0).text = name
    table.cell(i+1, 1).text = proj
    table.cell(i+1, 2).text = glob

doc.add_heading('4.7 Publishing to npm', level=2)
doc.add_paragraph('To publish a new version:')
steps = [
    'cd skills/',
    'npm version patch  (or minor/major)',
    'npm publish --access public',
    'This runs prepublishOnly which builds via obuild',
    'The compiled dist/ is included in the npm package',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════════
# 5. WEBSITE
# ══════════════════════════════════════════════
doc.add_heading('5. Website (website/)', level=1)

doc.add_heading('5.1 Tech Stack', level=2)
stack = [
    ('Framework', 'Next.js 16.x (App Router, React 19, Server Components)'),
    ('Styling', 'Tailwind CSS 4, shadcn/ui components, class-variance-authority'),
    ('UI Primitives', 'Radix UI (Select, Tabs, Slot)'),
    ('Markdown', 'react-markdown + remark-gfm + rehype-raw + rehype-sanitize'),
    ('Icons', 'Lucide React'),
    ('Animation', 'CSS animations (marquee, fade-up), IntersectionObserver'),
    ('Fonts', 'Inter (body), Bebas Neue (display/hero)'),
    ('Data', 'Filesystem (skills/index.json) + in-memory store'),
]
table = doc.add_table(rows=len(stack)+1, cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = 'Layer'
table.cell(0, 1).text = 'Technology'
table.cell(0, 0).paragraphs[0].runs[0].bold = True
table.cell(0, 1).paragraphs[0].runs[0].bold = True
for i, (layer, tech) in enumerate(stack):
    table.cell(i+1, 0).text = layer
    table.cell(i+1, 1).text = tech

doc.add_heading('5.2 Pages & Routing', level=2)
routes = [
    ('/', 'Homepage', 'Hero section + skills leaderboard table. Server component that fetches skills from index.json and merges with telemetry data. Renders SkillsList with search and category filters.'),
    ('/skills/[slug]', 'Skill Overview', 'Shows skill overview from overview.md or parsed description. Fetches from GitHub for community-installed skills.'),
    ('/skills/[slug]/documentation', 'Documentation', 'Full SKILL.md content rendered as markdown.'),
    ('/skills/[slug]/playbook', 'Playbook', 'Implementation guide (playbook.md).'),
    ('/skills/[slug]/advanced-patterns', 'Advanced', 'Advanced patterns (advanced-patterns.md).'),
    ('/skills/[slug]/cloud-integration', 'Cloud', 'Cloud execution guide (cloud-integration.md).'),
]
for route, name, desc in routes:
    p = doc.add_paragraph()
    run = p.add_run(f'{route} — {name}')
    run.bold = True
    doc.add_paragraph(desc)

doc.add_heading('5.3 API Endpoints', level=2)

apis = [
    ('GET /api/search', 'q (string), limit (number)',
     'Searches the in-memory skills store. Returns skills matching the query, '
     'sorted by install count. Used by the CLI "find" command and the homepage.'),
    ('GET /api/t', 'event, source, skills, agents, v, ci, global',
     'Telemetry ingestion endpoint. Receives install/remove/update events from the CLI. '
     'Stores events in memory and updates the skills search index. Returns 204 No Content.'),
    ('GET /api/github-stats', '(none)',
     'Returns stars and forks for the harishrajora/skills GitHub repo. '
     'Cached for 1 hour via Next.js revalidate.'),
    ('GET /api/audit', 'source, skills',
     'Security audit stub. Returns empty {}. In production, would return '
     'risk assessments for skills.'),
    ('GET /api/download/[owner]/[repo]/[slug]', '(path params)',
     'Download API stub. Returns empty files array. In production, would return '
     'pre-cached skill file snapshots for blob-based fast install.'),
]
for endpoint, params, desc in apis:
    p = doc.add_paragraph()
    run = p.add_run(endpoint)
    run.bold = True
    doc.add_paragraph(f'Parameters: {params}')
    doc.add_paragraph(desc)

doc.add_heading('5.4 Component Architecture', level=2)

components = [
    ('Header.tsx', 'Main navbar. Scroll-aware — uses relative positioning (scrolls with page). Shows logo, site name, GitHub stats, and navigation links. Different styles for homepage vs sub-pages.'),
    ('HeroTitle.tsx', 'Wraps the hero headline with a parallax scroll effect and delayed color animation (Stripe.dev-inspired). Uses useRef + scroll listener for translate3d transform.'),
    ('HeroCtaRow.tsx', 'Two CTA buttons: "Get started" (scrolls to install command) and "Browse skills" (scrolls to skills section). Uses smooth scrolling via element.scrollIntoView().'),
    ('AgentsMarquee.tsx', 'Infinite horizontal scroll of 12 agent logos (Claude, Cursor, Copilot, etc.). CSS animation with 30s duration, pauses on hover. Logos are grayscale with color on hover.'),
    ('InstallCommandBlock.tsx', 'Copy-to-clipboard component. In "hero" variant, shows the main install command. In "single"/"all" variants, shows per-agent git clone commands. The hero command is: npx agentskillshub-cli add https://github.com/harishrajora/skills.git'),
    ('SkillsList.tsx', 'The main skills listing — a ranked leaderboard table (inspired by skills.sh). Features: search input with "/" prefix, pill-style category filters with counts, numbered rows with skill logo/name/source/category. Client component with useMemo filtering.'),
    ('SkillLogo.tsx', 'Renders a framework icon from /public/skill-icons/. Maps skill paths to icon slugs via skillIcons.ts. Shows a fallback beaker icon when no icon exists.'),
    ('SkillDocNav.tsx', 'Tab navigation for skill documentation pages (Overview, Documentation, Playbook, etc.). Uses Radix UI Tabs with "pills" variant. Active tab determined by current pathname.'),
    ('SkillDocView.tsx', 'Renders markdown using ReactMarkdown with remark-gfm (tables, strikethrough), rehype-raw (HTML in markdown), and rehype-sanitize (XSS prevention). Styled with Tailwind prose classes.'),
    ('SkillOverviewDocView.tsx', 'Enhanced markdown renderer for overview pages. Custom components for tables (with impact badges), lists (with code highlighting for rule IDs), and headings (with border separators).'),
    ('GitHubStatsInline.tsx', 'Fetches and displays stars/forks for a specific GitHub repo. Used on skill detail pages — each skill shows stats from its own source repo, not a global count.'),
    ('CopyableUrlBlock.tsx', 'A URL/command display with one-click copy. Shows a code block with the install command and a copy button that shows a checkmark on success.'),
    ('ScrollReveal.tsx', 'IntersectionObserver component that adds "is-visible" class to elements with "animate-on-scroll" class, triggering CSS fade-up animations.'),
]
for name, desc in components:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    doc.add_paragraph(desc)

doc.add_heading('5.5 Library Files', level=2)

libs = [
    ('skills.ts', 'Core skill data loading. Functions:\n'
     '• getSkills() — loads from index.json + merges telemetry-installed skills\n'
     '• getSkillBySlug() — resolves URL slug to skill object\n'
     '• getSkillMarkdown() — reads local markdown or fetches from GitHub\n'
     '• getSkillInstallCommand() — generates npx install command\n'
     '• getGitHubSkillUrl() — builds GitHub browse URL\n'
     '• fetchInstalledSkills() — fetches from /api/search to get telemetry skills\n'
     '• fetchSkillMdFromGitHub() — fetches SKILL.md from GitHub for community skills\n'
     '• getSkillReadmeSections() — parses skill into structured sections'),
    ('skillsFs.ts', 'Server-only filesystem operations. Reads skills/index.json (with caching), '
     'SKILL.md files, and optional doc files (playbook.md, advanced-patterns.md, etc.) from the '
     'local skills/ directory.'),
    ('skillsHelpers.ts', 'Pure utility functions:\n'
     '• getSkillDisplayName() — converts path to human name with overrides\n'
     '• isHotSkill() — checks if skill is in the "hot" list\n'
     '• sortSkillsWithHotFirst() — sorts hot first, then alphabetical\n'
     '• toSkillSlug() — converts path to URL slug\n'
     '• categoryToSentenceCase() — "e2e-testing" → "E2e testing"'),
    ('skillsMapping.ts', 'Maps internal skill paths (e.g., "selenium-automation-skill") to '
     'LambdaTest/agent-skills folder names (e.g., "selenium-skill"). Used for GitHub URLs '
     'and install commands for the 45 pre-loaded skills.'),
    ('skillIcons.ts', 'Maps skill paths to SVG icon filenames in /public/skill-icons/. '
     'Each skill has a specific icon (e.g., selenium → selenium.svg, jest → jest.svg). '
     'Includes a cache-bust version parameter.'),
    ('store.ts', 'In-memory telemetry store. Two data structures:\n'
     '• events[] — array of all telemetry events\n'
     '• skillsMap — Map<string, SkillRecord> keyed by "source:skillName"\n'
     'addEvent() processes install events and updates skill records with install counts. '
     'searchSkills() filters and sorts by install count. '
     'getAllInstalledSkills() returns all tracked skills.\n\n'
     'IMPORTANT: This is in-memory only — data is lost on server restart. '
     'For production, replace with a database (Vercel Postgres, etc.).'),
    ('github.ts', 'Fetches live GitHub repo stats (stars, forks) from the GitHub API. '
     'Cached for 1 hour via Next.js revalidate. Used by the navbar stats display.'),
    ('types.ts', 'TypeScript interfaces: Skill, SkillsIndex, SkillFiles, DocType. '
     'The Skill interface includes an optional sourceRepo field for community-installed skills.'),
    ('parseDescription.ts', 'Parses skill description strings into structured data: '
     'whatItDoes, whenToUse[], triggers[]. Extracts "Use when..." and "Triggers on:" patterns.'),
    ('parseSkillMarkdown.ts', 'Splits markdown content by ## and ### headings into sections. '
     'Used to extract overview, samples, and advanced content from skill documentation.'),
    ('install.ts', 'Defines agent paths (Cursor, Claude, Copilot, Gemini, Codex) and generates '
     'git clone + cp commands for per-agent installation display.'),
]
for name, desc in libs:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    doc.add_paragraph(desc)

doc.add_heading('5.6 Skills Data & Catalog', level=2)
doc.add_paragraph(
    'The website ships with 45 pre-loaded test automation skills in the skills/ directory. '
    'Each skill has its own folder containing:'
)
doc.add_paragraph('SKILL.md — Main skill definition with YAML frontmatter (name, description)', style='List Bullet')
doc.add_paragraph('overview.md — Catalog-derived overview content', style='List Bullet')
doc.add_paragraph('playbook.md — Implementation guide with code samples', style='List Bullet')
doc.add_paragraph('advanced-patterns.md — Advanced usage patterns', style='List Bullet')
doc.add_paragraph('cloud-integration.md — Cloud execution guide', style='List Bullet')

doc.add_paragraph(
    'The skills/index.json file is the master catalog. It contains metadata for all 45 skills: '
    'name, path, description, languages, category, and file paths. Categories include: '
    'e2e-testing, unit-testing, bdd-testing, mobile-testing, cloud-testing, devops, visual-testing.'
)

doc.add_heading('5.7 Telemetry-Driven Discovery', level=2)
doc.add_paragraph(
    'When a user installs any skill from any GitHub repo via the CLI, the telemetry event '
    'reaches the website\'s /api/t endpoint. This adds the skill to the in-memory store. '
    'The homepage\'s getSkills() function then fetches from /api/search and merges these '
    'community-installed skills with the static catalog. This means:'
)
doc.add_paragraph('Any skill installed via the CLI automatically appears on the homepage', style='List Bullet')
doc.add_paragraph('Each community skill links back to its actual GitHub source repo', style='List Bullet')
doc.add_paragraph('Documentation is fetched live from the source repo\'s SKILL.md or README.md', style='List Bullet')
doc.add_paragraph('Install counts are tracked and used for search ranking', style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════════════════
# 6. END-TO-END DATA FLOW
# ══════════════════════════════════════════════
doc.add_heading('6. Data Flow: End-to-End', level=1)

doc.add_heading('6.1 Installing a Skill', level=2)
doc.add_paragraph(
    'Example: npx agentskillshub-cli add https://github.com/vercel-labs/agent-browser'
)
steps = [
    'CLI parses URL → {url: "https://github.com/vercel-labs/agent-browser.git", owner: "vercel-labs", repo: "agent-browser"}',
    'CLI tries blob install → fails (vercel-labs not in allowlist)',
    'CLI clones repo → git clone --depth 1 to temp dir',
    'CLI finds SKILL.md → discovers "agent-browser" skill with name "Agent Browser"',
    'CLI detects agents → finds Claude Code (.claude/), Cursor (.cursor/)',
    'CLI installs → symlinks skill files to .claude/skills/agent-browser/ and .cursor/skills/agent-browser/',
    'CLI sends telemetry → GET /api/t?event=install&source=vercel-labs/agent-browser&skills=agent-browser',
    'Website /api/t → stores {name: "agent-browser", source: "vercel-labs/agent-browser", installCount: 1}',
    'Next homepage load → getSkills() fetches /api/search → merges "agent-browser" into listing',
    'User visits /skills/agent-browser → website fetches SKILL.md from GitHub → renders documentation',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_heading('6.2 Searching for Skills', level=2)
doc.add_paragraph(
    'Example: npx agentskillshub-cli find selenium'
)
steps = [
    'CLI sends GET /api/search?q=selenium&limit=10 to the website',
    'Website searches in-memory store for skills matching "selenium"',
    'Returns results sorted by install count',
    'CLI displays results in interactive fzf-style UI',
    'User selects a skill → CLI runs "add" with that skill',
]
for i, step in enumerate(steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_page_break()

# ══════════════════════════════════════════════
# 7. ENVIRONMENT VARIABLES
# ══════════════════════════════════════════════
doc.add_heading('7. Environment Variables', level=1)

envs = [
    ('SKILLS_API_URL', 'CLI', 'Override the website URL for search, telemetry, and download APIs. Default: https://www.agentskillsforall.com. Use http://localhost:3000 for local development.'),
    ('SKILLS_DOWNLOAD_URL', 'CLI', 'Override the download API URL. Defaults to SKILLS_API_URL value.'),
    ('DISABLE_TELEMETRY', 'CLI', 'Set to 1 to disable all telemetry.'),
    ('DO_NOT_TRACK', 'CLI', 'Alternative telemetry disable flag.'),
    ('INSTALL_INTERNAL_SKILLS', 'CLI', 'Set to 1 to show internal/WIP skills in search.'),
    ('NEXT_PUBLIC_SITE_URL', 'Website', 'The website\'s own URL. Used for internal API calls. Default: http://localhost:3000.'),
]
table = doc.add_table(rows=len(envs)+1, cols=3)
table.style = 'Light Grid Accent 1'
for j, h in enumerate(['Variable', 'Used By', 'Description']):
    table.cell(0, j).text = h
    table.cell(0, j).paragraphs[0].runs[0].bold = True
for i, (var, used, desc) in enumerate(envs):
    table.cell(i+1, 0).text = var
    table.cell(i+1, 1).text = used
    table.cell(i+1, 2).text = desc

# ══════════════════════════════════════════════
# 8. DEPLOYMENT GUIDE
# ══════════════════════════════════════════════
doc.add_heading('8. Deployment Guide', level=1)

doc.add_heading('8.1 CLI (npm)', level=2)
doc.add_paragraph('The CLI is published to npm. To release a new version:')
code = 'cd skills/\nnpm version patch   # or minor/major\nnpm publish --access public'
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('8.2 Website (Vercel)', level=2)
doc.add_paragraph('The website is deployed on Vercel. Configuration:')
doc.add_paragraph('Repository: github.com/harishrajora/agentskillsforall', style='List Bullet')
doc.add_paragraph('Root Directory: website (set in Vercel dashboard)', style='List Bullet')
doc.add_paragraph('Framework: Next.js', style='List Bullet')
doc.add_paragraph('Build Command: next build (or npm run build)', style='List Bullet')
doc.add_paragraph('The root package.json includes "next" as a dependency for Vercel\'s framework detection.', style='List Bullet')
doc.add_paragraph(
    'After deployment, update the custom domain to agentskillsforall.com in Vercel\'s '
    'domain settings. The CLI will then send telemetry to the production website automatically.'
)

doc.add_heading('8.3 Connecting CLI to Website', level=2)
doc.add_paragraph(
    'The CLI\'s telemetry.ts sends events to https://www.agentskillsforall.com/api/t by default. '
    'Once the website is deployed to that domain, the loop is complete:\n\n'
    '1. User installs a skill via CLI → telemetry goes to production website\n'
    '2. Website stores the event → skill appears in search and homepage\n'
    '3. Another user searches via CLI → finds the skill → installs it\n\n'
    'For local development, use SKILLS_API_URL=http://localhost:3000 before the npx command.'
)

doc.add_page_break()

# ══════════════════════════════════════════════
# 9. DEVELOPMENT WORKFLOW
# ══════════════════════════════════════════════
doc.add_heading('9. Development Workflow', level=1)

doc.add_heading('9.1 CLI Development', level=2)
code = '# Install dependencies\ncd skills/\npnpm install\n\n# Run CLI in dev mode\nnode src/cli.ts add owner/repo\n\n# Build\nnpm run build\n\n# Test locally\nnode bin/cli.mjs --help\n\n# Run tests\nnpm test'
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('9.2 Website Development', level=2)
code = '# Install dependencies\ncd website/\nnpm install\n\n# Start dev server (http://localhost:3000)\nnpm run dev\n\n# Build for production\nnpm run build\n\n# Start production server\nnpm start'
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('9.3 Testing End-to-End Locally', level=2)
code = '# Terminal 1: Start website\ncd website/ && npm run dev\n\n# Terminal 2: Install a skill with local telemetry\nSKILLS_API_URL=http://localhost:3000 npx agentskillshub-cli add owner/repo\n\n# Check homepage to see the skill appear\nopen http://localhost:3000'
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(10)

# ══════════════════════════════════════════════
# 10. KEY DESIGN DECISIONS
# ══════════════════════════════════════════════
doc.add_heading('10. Key Design Decisions', level=1)

decisions = [
    ('Symlinks over Copies', 'By default, skills are symlinked rather than copied to agent directories. This means updating the source updates all agents automatically. The --copy flag is available for environments that don\'t support symlinks.'),
    ('In-Memory Store', 'The website uses an in-memory store for telemetry data. This is simple and fast but data is lost on server restart. For production scale, this should be replaced with a database (Vercel Postgres, Supabase, etc.).'),
    ('GitHub Fallback for Documentation', 'Community-installed skills don\'t have local files. The website fetches SKILL.md directly from the source GitHub repo, with README.md as a final fallback. This means any skill installed via the CLI gets documentation automatically.'),
    ('Telemetry-Driven Discovery', 'Rather than requiring skill authors to register on a central registry, skills are discovered through usage telemetry. Every install event adds the skill to the searchable index. This is the same approach used by skills.sh.'),
    ('SKILLS_API_URL for Local Dev', 'The CLI uses an environment variable to override the API URL. This allows local development without deploying the website, while defaulting to the production URL for end users.'),
    ('Static + Dynamic Catalog', 'The homepage merges a static catalog (45 pre-loaded skills from index.json) with dynamic telemetry data. This ensures the page always has content even when the in-memory store is empty.'),
    ('Dark Theme Only', 'The website uses a dark theme exclusively (class="dark" on html). This matches the developer tool aesthetic and reduces visual complexity.'),
]
for title, desc in decisions:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(desc)

# ══════════════════════════════════════════════
# 11. TROUBLESHOOTING
# ══════════════════════════════════════════════
doc.add_heading('11. Troubleshooting', level=1)

issues = [
    ('Skills don\'t appear on homepage after install', 'The in-memory store resets on server restart. Re-install the skill or manually hit /api/t with the event data. For production, implement a persistent database.'),
    ('CLI sends telemetry to wrong URL', 'Check SKILLS_API_URL env var. The default is https://www.agentskillsforall.com. For local testing, set SKILLS_API_URL=http://localhost:3000.'),
    ('Vercel build fails with "No Next.js version detected"', 'Ensure the root package.json has "next" in dependencies. Set Root Directory to "website" in Vercel dashboard. Framework preset should be "Next.js".'),
    ('"npx agentskillshub-cli" runs old version', 'npm caches packages. Run "npx agentskillshub-cli@latest" to force the latest, or clear cache with "npx clear-npx-cache".'),
    ('Skill documentation shows raw HTML', 'The SkillDocView component uses rehype-raw and rehype-sanitize to handle HTML in markdown. If you see raw HTML, ensure these packages are installed in website/.'),
    ('GitHub stars/forks show wrong values', 'GitHubStatsInline.tsx fetches stats per skill from each skill\'s sourceRepo. Catalog skills default to harishrajora/skills. Community skills show their actual source repo stats.'),
    ('Blob install fails and falls back to clone', 'Blob install only works for allowlisted owners (LambdaTest, harishrajora). Other repos always use git clone. This is expected behavior.'),
]
for title, desc in issues:
    p = doc.add_paragraph()
    run = p.add_run(f'Issue: {title}')
    run.bold = True
    doc.add_paragraph(f'Solution: {desc}')

# ── Save ──
output_path = '/Users/harishr/Documents/repos/skills-collection/docs/AgentSkillsForAll_Engineering_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
